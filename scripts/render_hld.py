#!/usr/bin/env python
"""Render HLD_Document.md into the two files a reviewer actually opens:

    python scripts/render_hld.py

Writes HLD_Document.pdf and HLD_Document.docx beside the markdown, so the
markdown stays the single source of truth and the rendered copies cannot drift
from it. Needs reportlab, python-docx and Pillow only - no pandoc, no
libreoffice. Run this last, after every prose edit is final.
"""
import html
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "HLD_Document.md"

INK = "#1b2733"
MUTED = "#5b6b7a"
RULE = "#9fb0c0"
HEAD_FILL = "#e8f1f8"
CODE_FILL = "#f4f7fa"


def _align(cell):
    if cell.startswith(":") and cell.endswith(":"):
        return "center"
    if cell.endswith(":"):
        return "right"
    return "left"


def parse(md):
    """Markdown -> a flat list of (kind, payload). Only the subset the HLD uses:
    headings, paragraphs, fenced code, pipe tables, lists and images."""
    lines = md.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(buf)))
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            blocks.append((f"h{len(heading.group(1))}", heading.group(2).strip()))
            i += 1
            continue

        image = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image:
            blocks.append(("image", (image.group(2), image.group(1))))
            i += 1
            continue

        if line.startswith("|"):
            rows, aligns = [], None
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                i += 1
                if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    aligns = [_align(c) for c in cells]
                    continue
                rows.append(cells)
            blocks.append(("table", (rows, aligns)))
            continue

        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
                # A wrapped bullet continues on an indented line.
                while (i < len(lines) and lines[i].startswith("  ")
                       and lines[i].strip()
                       and not re.match(r"^[-*]\s+", lines[i].strip())):
                    items[-1] += " " + lines[i].strip()
                    i += 1
            blocks.append(("bullets", items))
            continue

        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ordered", items))
            continue

        # Paragraph. A line ending in two spaces is a hard break, which is how the
        # document's author/version/scenario header is written.
        buf = [line]
        hard = raw.endswith("  ")
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if nxt.strip()[0] in "|#" or nxt.strip().startswith(("```", "![", "- ", "* ")):
                break
            buf.append(("\n" if hard else " ") + nxt.strip())
            hard = nxt.endswith("  ")
            i += 1
        blocks.append(("para", "".join(buf)))
    return blocks


def inline_pdf(text):
    out = html.escape(text, quote=False)
    out = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", out)
    out = re.sub(r"`(.+?)`", r'<font face="Courier" size="8">\1</font>', out)
    return out


def code_pdf(text):
    out = html.escape(text, quote=False).replace(" ", "&nbsp;")
    return out.replace("\n", "<br/>")


def render_pdf(blocks, path):
    from PIL import Image as PILImage
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (Image, ListFlowable, ListItem, Paragraph,
                                    SimpleDocTemplate, Spacer, Table, TableStyle)

    base = getSampleStyleSheet()["BodyText"]
    body = ParagraphStyle("body", parent=base, fontName="Helvetica", fontSize=9.5,
                          leading=13.6, textColor=colors.HexColor(INK), spaceAfter=6)
    styles = {
        "h1": ParagraphStyle("h1", parent=body, fontName="Helvetica-Bold", fontSize=17,
                             leading=21, spaceBefore=0, spaceAfter=10),
        "h2": ParagraphStyle("h2", parent=body, fontName="Helvetica-Bold", fontSize=12.5,
                             leading=16, spaceBefore=15, spaceAfter=6),
        "h3": ParagraphStyle("h3", parent=body, fontName="Helvetica-Bold", fontSize=10.5,
                             leading=14, spaceBefore=11, spaceAfter=5,
                             textColor=colors.HexColor(MUTED)),
    }
    styles["h4"] = styles["h3"]
    mono = ParagraphStyle("mono", parent=body, fontName="Courier", fontSize=8,
                          leading=10.6, spaceAfter=0)
    cell = ParagraphStyle("cell", parent=body, fontSize=8.2, leading=11, spaceAfter=0)
    cellhead = ParagraphStyle("cellhead", parent=cell, fontName="Helvetica-Bold")
    caption = ParagraphStyle("caption", parent=body, fontSize=8, spaceBefore=3,
                             textColor=colors.HexColor(MUTED))

    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=1.9 * cm, rightMargin=1.9 * cm,
                            topMargin=1.7 * cm, bottomMargin=1.7 * cm,
                            title="Kapture Finance - Maya Collections Voicebot",
                            author="Sachin Eldho")
    width = doc.width
    story = []

    for kind, payload in blocks:
        if kind in styles:
            story.append(Paragraph(inline_pdf(payload), styles[kind]))
        elif kind == "para":
            story.append(Paragraph(inline_pdf(payload).replace("\n", "<br/>"), body))
        elif kind in ("bullets", "ordered"):
            items = [ListItem(Paragraph(inline_pdf(t), body), leftIndent=15)
                     for t in payload]
            story.append(ListFlowable(
                items, leftIndent=15, bulletFontSize=8,
                bulletType="1" if kind == "ordered" else "bullet"))
            story.append(Spacer(1, 6))
        elif kind == "code":
            story.append(Table(
                [[Paragraph(code_pdf(payload), mono)]], colWidths=[width],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(CODE_FILL)),
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(RULE)),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6)])))
            story.append(Spacer(1, 9))
        elif kind == "table":
            rows, aligns = payload
            if not rows:
                continue
            ncols = max(len(row) for row in rows)
            rows = [row + [""] * (ncols - len(row)) for row in rows]
            # Column widths track the longest cell, clamped so one verbose column
            # cannot squeeze the rest to unreadable slivers - but never narrower
            # than the longest unbreakable token, or reportlab splits mid-word.
            weights = []
            for c in range(ncols):
                longest_cell = max(len(row[c]) for row in rows)
                longest_word = max((len(word) for row in rows
                                    for word in row[c].split()), default=0)
                weights.append(max(7, longest_word + 2, min(longest_cell, 44)))
            total = sum(weights)
            data = [[Paragraph(inline_pdf(c), cellhead) for c in rows[0]]]
            data += [[Paragraph(inline_pdf(c), cell) for c in row] for row in rows[1:]]
            style = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(HEAD_FILL)),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(RULE)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            for idx, align in enumerate((aligns or [])[:ncols]):
                if align != "left":
                    style.append(("ALIGN", (idx, 0), (idx, -1), align.upper()))
            story.append(Table(data, colWidths=[width * w / total for w in weights],
                               style=TableStyle(style), repeatRows=1))
            story.append(Spacer(1, 9))
        elif kind == "image":
            src, alt = payload
            file = ROOT / src
            if not file.exists():
                print(f"  skipped missing image {src}")
                continue
            with PILImage.open(file) as img:
                iw, ih = img.size
            scale = min(width / iw, 1.0)
            story.append(Image(str(file), iw * scale, ih * scale))
            if alt:
                story.append(Paragraph(f"<i>{html.escape(alt)}</i>", caption))
            story.append(Spacer(1, 9))
    doc.build(story)
    print(f"wrote {path.name}")


def add_runs(paragraph, text, size=None):
    """Inline markdown -> docx runs. Bold and code spans only; the HLD uses no
    other inline markup."""
    from docx.shared import Pt

    for line_index, segment in enumerate(text.split("\n")):
        if line_index:
            paragraph.add_run().add_break()
        for chunk in re.split(r"(\*\*.+?\*\*|`[^`]+`)", segment):
            if not chunk:
                continue
            if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
                run = paragraph.add_run(chunk[2:-2])
                run.bold = True
            elif chunk.startswith("`") and chunk.endswith("`") and len(chunk) > 2:
                run = paragraph.add_run(chunk[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(size - 0.5 if size else 9)
            else:
                run = paragraph.add_run(chunk)
            if size:
                run.font.size = run.font.size or Pt(size)


def render_docx(blocks, path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    body_width = Inches(6.3)

    for kind, payload in blocks:
        if kind == "h1":
            heading = doc.add_heading(level=0)
            add_runs(heading, payload)
        elif kind in ("h2", "h3", "h4"):
            heading = doc.add_heading(level=2 if kind == "h2" else 3)
            add_runs(heading, payload)
        elif kind == "para":
            add_runs(doc.add_paragraph(), payload)
        elif kind in ("bullets", "ordered"):
            style = "List Number" if kind == "ordered" else "List Bullet"
            for item in payload:
                add_runs(doc.add_paragraph(style=style), item)
        elif kind == "code":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            run = paragraph.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x1B, 0x27, 0x33)
        elif kind == "table":
            rows, _ = payload
            if not rows:
                continue
            ncols = max(len(row) for row in rows)
            rows = [row + [""] * (ncols - len(row)) for row in rows]
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            table.autofit = True
            for r, row in enumerate(rows):
                for c, text in enumerate(row):
                    para = table.cell(r, c).paragraphs[0]
                    add_runs(para, text, size=8.5)
                    if r == 0:
                        for run in para.runs:
                            run.bold = True
        elif kind == "image":
            src, alt = payload
            file = ROOT / src
            if not file.exists():
                print(f"  skipped missing image {src}")
                continue
            doc.add_picture(str(file), width=body_width)
            if alt:
                caption = doc.add_paragraph()
                run = caption.add_run(alt)
                run.italic = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7A)
    doc.save(str(path))
    print(f"wrote {path.name}")


if __name__ == "__main__":
    blocks = parse(SRC.read_text(encoding="utf-8"))
    counts = {}
    for kind, _ in blocks:
        counts[kind] = counts.get(kind, 0) + 1
    print(f"parsed {len(blocks)} blocks: " +
          ", ".join(f"{k}={v}" for k, v in sorted(counts.items())))
    render_pdf(blocks, ROOT / "HLD_Document.pdf")
    render_docx(blocks, ROOT / "HLD_Document.docx")
